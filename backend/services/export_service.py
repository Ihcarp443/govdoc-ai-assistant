# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas


# class ExportService:

#     def __init__(self):
#         pass

#     def create_docx(self,report_text: str,file_name: str = "report.docx") -> str:

#         doc = Document()

#         for line in report_text.split("\n"):

#             line = line.strip()
#             if not line:
#                 continue

#             if ":" in line:
#                 title, content = line.split(":", 1)
#                 doc.add_heading(title.strip(), level=1)
#                 doc.add_paragraph(content.strip())
#             else:

#                 doc.add_paragraph(line)

#         doc.save(file_name)
#         return file_name

#     def create_pdf(self,report_text: str,file_name: str = "report.pdf") -> str:

#         pdf = canvas.Canvas(file_name,pagesize=letter)
#         width, height = letter

#         y = height - 50

#         for line in report_text.split("\n"):
#             if y < 50:
#                 pdf.showPage()
#                 y = height - 50

#             pdf.drawString(50, y, line)
#             y -= 18

#         pdf.save()
#         return file_name

#     def export(self,text: str,file_type: str,file_name: str):

#         if file_type.lower() == "pdf":
#             return self.create_pdf(text, f"{file_name}.pdf")

#         elif file_type.lower() == "docx":
#             return self.create_docx(text, f"{file_name}.docx")

#         raise ValueError("Unsupported format")


import os
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

from markdown import markdown
from bs4 import BeautifulSoup


class ExportService:

    def __init__(self):
        os.makedirs("exports", exist_ok=True)

    # ----------------------------
    # PUBLIC
    # ----------------------------

    def export(self, markdown, file_type, file_name, path):

       

        os.makedirs(os.path.dirname(path), exist_ok=True)

        if file_type == "docx":
            self.export_docx(markdown, path)

        elif file_type == "pdf":
            self.export_pdf(markdown, path)

        else:
            raise ValueError("Unsupported format")

        return path

    # ----------------------------
    # DOCX
    # ----------------------------

    def export_docx(self, markdown_text, path):

        html = markdown(
            markdown_text,
            extensions=["tables"]
        )

        soup = BeautifulSoup(html, "html.parser")

        doc = Document()

        for element in soup.children:

            if element.name == "h1":
                p = doc.add_heading(level=1)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(element.get_text())

            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)

            elif element.name == "h3":
                doc.add_heading(element.get_text(), level=3)

            elif element.name == "p":
                self._add_runs(doc.add_paragraph(), element)

            elif element.name == "ul":

                for li in element.find_all("li", recursive=False):
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_runs(p, li)

            elif element.name == "table":

                rows = element.find_all("tr")

                cols = rows[0].find_all(["th", "td"])

                table = doc.add_table(
                    rows=len(rows),
                    cols=len(cols)
                )

                table.style = "Table Grid"

                for i, row in enumerate(rows):

                    cells = row.find_all(["th", "td"])

                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()

        doc.save(path)

    # ----------------------------
    # PDF
    # ----------------------------

    def export_pdf(self, markdown_text, path):

        html = markdown(
            markdown_text,
            extensions=["tables"]
        )

        soup = BeautifulSoup(html, "html.parser")

        styles = getSampleStyleSheet()

        story = []

        for element in soup.children:

            if element.name == "h1":

                story.append(
                    Paragraph(
                        f"<b>{element.get_text()}</b>",
                        styles["Title"]
                    )
                )

                story.append(Spacer(1, 12))

            elif element.name == "h2":

                story.append(
                    Paragraph(
                        f"<b>{element.get_text()}</b>",
                        styles["Heading2"]
                    )
                )

            elif element.name == "p":

                story.append(
                    Paragraph(
                        str(element),
                        styles["BodyText"]
                    )
                )

                story.append(Spacer(1, 8))

            elif element.name == "ul":

                for li in element.find_all("li", recursive=False):

                    story.append(
                        Paragraph(
                            "• " + li.get_text(),
                            styles["BodyText"]
                        )
                    )

            elif element.name == "table":

                data = []

                for row in element.find_all("tr"):

                    data.append([
                        cell.get_text()
                        for cell in row.find_all(["th", "td"])
                    ])

                table = Table(data)

                table.setStyle(TableStyle([
                    ("GRID", (0,0), (-1,-1), 1, colors.black),
                    ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
                    ("BOTTOMPADDING",(0,0),(-1,0),8),
                    ("ALIGN",(0,0),(-1,-1),"CENTER"),
                ]))

                story.append(table)
                story.append(Spacer(1, 12))

        pdf = SimpleDocTemplate(path)

        pdf.build(story)

    # ----------------------------
    # Helpers
    # ----------------------------

    def _add_runs(self, paragraph, element):

        for child in element.children:

            if getattr(child, "name", None) == "strong":
                run = paragraph.add_run(child.get_text())
                run.bold = True

            elif getattr(child, "name", None) == "em":
                run = paragraph.add_run(child.get_text())
                run.italic = True

            else:
                paragraph.add_run(str(child))